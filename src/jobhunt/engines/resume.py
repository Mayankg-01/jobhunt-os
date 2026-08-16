from __future__ import annotations

import hashlib
import json
from datetime import datetime
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from jobhunt.core.config import settings
from jobhunt.utils.ai import generate_resume_id, quality_gate_resume, tailor_resume


class ResumeParser:
    @staticmethod
    def parse_pdf(file_path: Path) -> dict:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return ResumeParser._parse_text(text)

    @staticmethod
    def parse_docx(file_path: Path) -> dict:
        doc = DocxDocument(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return ResumeParser._parse_text(text)

    @staticmethod
    def _parse_text(text: str) -> dict:
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        return {"raw_text": text, "lines": lines, "parsed_at": datetime.utcnow().isoformat()}


class ResumeTailor:
    def __init__(self):
        self.resumes_dir = settings.data_dir / "resumes"
        self.resumes_dir.mkdir(parents=True, exist_ok=True)

    async def tailor(
        self,
        base_resume_path: Path,
        job_description: str,
        user_profile: dict,
        output_dir: Path | None = None,
    ) -> dict:
        base_resume = ResumeParser.parse_pdf(base_resume_path) if base_resume_path.suffix == ".pdf" else ResumeParser.parse_docx(base_resume_path)

        tailored = await tailor_resume(base_resume, job_description, user_profile)
        if "error" in tailored:
            return tailored

        quality = await quality_gate_resume(tailored, job_description)

        resume_id = generate_resume_id()
        if output_dir is None:
            output_dir = self.resumes_dir / resume_id
        output_dir.mkdir(parents=True, exist_ok=True)

        json_path = output_dir / "resume.json"
        with open(json_path, "w") as f:
            json.dump(tailored, f, indent=2)

        docx_path = output_dir / "resume.docx"
        self._generate_docx(tailored, docx_path)

        pdf_path = output_dir / "resume.pdf"
        self._generate_pdf(docx_path, pdf_path)

        content_hash = self._hash_file(pdf_path)

        return {
            "resume_id": resume_id,
            "resume": tailored,
            "quality_gate": quality,
            "files": {
                "json": str(json_path),
                "docx": str(docx_path),
                "pdf": str(pdf_path),
            },
            "content_hash": content_hash,
            "created_at": datetime.utcnow().isoformat(),
        }

    def _generate_docx(self, resume: dict, output_path: Path) -> None:
        doc = DocxDocument()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10.5)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        for section in ["topMargin", "bottomMargin", "leftMargin", "rightMargin"]:
            setattr(doc.sections[0], section, Inches(0.6))

        header = resume.get("header", {})
        name = header.get("name", "").upper()
        contact_parts = []
        for field in ["email", "phone", "location", "linkedin", "github", "portfolio"]:
            val = header.get(field, "")
            if val:
                contact_parts.append(val)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        p.space_after = Pt(2)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.space_after = Pt(8)

        self._add_section(doc, "PROFESSIONAL SUMMARY", resume.get("summary", ""))

        skills = resume.get("skills", {})
        if skills:
            skills_text = []
            for category, items in skills.items():
                if items:
                    skills_text.append(f"{category.title()}: {', '.join(items)}")
            self._add_section(doc, "TECHNICAL SKILLS", "\n".join(skills_text))

        experience = resume.get("experience", [])
        if experience:
            self._add_section_header(doc, "PROFESSIONAL EXPERIENCE")
            for exp in experience:
                p = doc.add_paragraph()
                run = p.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}")
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

                meta = []
                if exp.get("dates"):
                    meta.append(exp["dates"])
                if exp.get("location"):
                    meta.append(exp["location"])
                p = doc.add_paragraph(" | ".join(meta))
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                p.runs[0].italic = True
                p.space_after = Pt(2)

                for bullet in exp.get("bullets", []):
                    p = doc.add_paragraph(bullet, style="List Bullet")
                    p.runs[0].font.size = Pt(10)
                    p.space_after = Pt(1)
                p.space_after = Pt(6)

        education = resume.get("education", [])
        if education:
            self._add_section_header(doc, "EDUCATION")
            for edu in education:
                p = doc.add_paragraph()
                run = p.add_run(f"{edu.get('degree', '')}, {edu.get('school', '')}")
                run.bold = True
                run.font.size = Pt(10.5)
                if edu.get("year"):
                    p.add_run(f" | {edu['year']}").font.size = Pt(10)
                if edu.get("honors"):
                    p.add_run(f" | {edu['honors']}").font.size = Pt(10).italic = True

        projects = resume.get("projects", [])
        if projects:
            self._add_section_header(doc, "PROJECTS")
            for proj in projects:
                p = doc.add_paragraph()
                run = p.add_run(proj.get("name", ""))
                run.bold = True
                run.font.size = Pt(10.5)
                if proj.get("tech_stack"):
                    p.add_run(f" | {', '.join(proj['tech_stack'])}").font.size = Pt(9).font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                if proj.get("description"):
                    p = doc.add_paragraph(proj["description"])
                    p.runs[0].font.size = Pt(10)

        certs = resume.get("certifications", [])
        if certs:
            self._add_section(doc, "CERTIFICATIONS", "\n".join(certs))

        doc.save(output_path)

    def _add_section_header(self, doc: DocxDocument, title: str) -> None:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.name = "Calibri"
        p.space_before = Pt(8)
        p.space_after = Pt(4)

        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "1A1A2E",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_section(self, doc: DocxDocument, title: str, content: str) -> None:
        self._add_section_header(doc, title)
        p = doc.add_paragraph(content)
        p.runs[0].font.size = Pt(10)
        p.space_after = Pt(6)

    def _generate_pdf(self, docx_path: Path, pdf_path: Path) -> None:
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(pdf_path))
        except Exception:
            import subprocess
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", str(pdf_path.parent), str(docx_path)
            ], check=False, capture_output=True)

    def _hash_file(self, file_path: Path) -> str:
        hasher = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                hasher.update(chunk)
        return hasher.hexdigest()


class ResumeQualityGate:
    MIN_SCORE = 7
    HARD_FAIL_MIN_BULLETS = 3

    @staticmethod
    def check(resume: dict, job_description: str) -> dict:
        issues = []
        warnings = []

        experience = resume.get("experience", [])
        total_bullets = sum(len(exp.get("bullets", [])) for exp in experience)
        if total_bullets < ResumeQualityGate.HARD_FAIL_MIN_BULLETS:
            issues.append(f"Too few bullets: {total_bullets} (minimum {ResumeQualityGate.HARD_FAIL_MIN_BULLETS})")

        skills = resume.get("skills", {})
        if not any(skills.values()):
            warnings.append("No skills listed")

        if not resume.get("header", {}).get("email"):
            issues.append("Missing email in header")

        if not resume.get("summary"):
            warnings.append("Missing professional summary")

        jd_keywords = set(job_description.lower().split())
        resume_text = json.dumps(resume).lower()
        matched = sum(1 for kw in jd_keywords if len(kw) > 3 and kw in resume_text)
        keyword_ratio = matched / max(len(jd_keywords), 1)
        if keyword_ratio < 0.15:
            warnings.append(f"Low JD keyword match: {keyword_ratio:.1%}")

        return {
            "checks": {
                "bullets": total_bullets >= ResumeQualityGate.HARD_FAIL_MIN_BULLETS,
                "skills": bool(any(skills.values())),
                "contact": bool(resume.get("header", {}).get("email")),
                "summary": bool(resume.get("summary")),
                "keywords": keyword_ratio >= 0.15,
            },
            "issues": issues,
            "warnings": warnings,
            "passed": len(issues) == 0,
        }
